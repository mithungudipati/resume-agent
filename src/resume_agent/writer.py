"""Markdown to DOCX converter with professional formatting."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Navy color for headings
NAVY = RGBColor(0x00, 0x2B, 0x5C)
BLACK = RGBColor(0x00, 0x00, 0x00)


def write_resume_docx(markdown_text: str, output_path: str) -> str:
    """Convert Markdown resume text to a professionally formatted DOCX.

    Returns the output file path.
    """
    doc = Document()

    # Set margins to 0.75 inches
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = BLACK

    lines = markdown_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            # H3 — Job title / sub-heading
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            _add_formatted_text(p, text, bold=True, size=Pt(12), color=BLACK)

        elif stripped.startswith("## "):
            # H2 — Section heading
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.space_before = Pt(10)
            _add_formatted_text(p, text, bold=True, size=Pt(13), color=NAVY)
            # Add a bottom border via a thin line
            _add_bottom_border(p)

        elif stripped.startswith("# "):
            # H1 — Name
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_formatted_text(p, text, bold=True, size=Pt(18), color=NAVY)

        elif stripped.startswith("- ") or stripped.startswith("* "):
            # Bullet point
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, text, size=Pt(11))

        else:
            # Regular paragraph
            p = doc.add_paragraph()
            _add_inline_formatting(p, stripped, size=Pt(11))

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def _add_formatted_text(paragraph, text, bold=False, size=None, color=None):
    """Add a run with specific formatting to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    run.font.name = "Calibri"


def _add_inline_formatting(paragraph, text, size=None):
    """Parse inline **bold** markers and add runs accordingly."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            run.bold = False
        run.font.name = "Calibri"
        if size:
            run.font.size = size


def _add_bottom_border(paragraph):
    """Add a thin bottom border to a paragraph (section divider)."""
    from docx.oxml.ns import qn
    from lxml import etree

    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "002B5C")
